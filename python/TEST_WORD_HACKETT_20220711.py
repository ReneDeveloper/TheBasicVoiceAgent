
#pip install python-docx 
#
#https://pypi.org/project/docx2pdf/


#https://www.geeksforgeeks.org/read-json-file-using-python/

#--------------------------------------------------------


#brightness_4
# Python program to read 
# json file 
  
  
import json 
  

import docx
 
from docx.shared import Inches


def read_template_WORD(doc_template):
    read_template_WORD_OUT = None
    read_template_WORD_OUT = docx.Document(doc_template)
    return read_template_WORD_OUT

def add_PART(doc_template, DIC_PART):

    DIC_PART = {}
    DIC_PART["PAR1"] = 'A plain paragraph having some '
    DIC_PART["HEAD1"] = 'TITULO NIVEL 1'
    DIC_PART["HEAD2"] = 'TITULO NIVEL 2'

    p = doc_template.add_paragraph(DIC_PART["PAR1"])
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    doc_template.add_heading(DIC_PART["HEAD1"], level=1)
    doc_template.add_heading(DIC_PART["HEAD2"], level=2)

    doc_template.add_paragraph('Intense quote', style='Intense Quote')

    doc_template.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    doc_template.add_paragraph(
        'first item in ordered list', style='List Number'
    )
    return doc_template



def add_TABLE(doc_template, DIC_TABLE):




    # JSON file 
    f = open ('PARAM_WORD.json', "r") 
    
    # Reading from file 
    data = json.loads(f.read()) 
      
    # Iterating through the json 
    # list 
    for fila in data:
        #print(fila)
        print(f"fila[Fila]:{fila['Fila']}") 
        print(f"fila[Cantidad]:{fila['Cantidad']}") 
      
    # Closing file 
    f.close() 

    DIC_TABLE = (
        (3, '000', '0000'),
        (7, '000', '00000'),
        (4, '000', '00000000, spam, eggs, and spam')
    )

    table = doc_template.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in DIC_TABLE:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    doc_template.add_page_break()

    return doc_template

def main__WORD():

    # JSON string 
    WORD_CONFIG = '{"Title": "Bob", "languages": "English"}'
    
    # deserializes into dict  
    # and returns dict. 
    y = json.loads(a) 
      
    print("JSON string = ", y) 

    


      








    #---------------------------------------------------------
    """
    #pip install python-docx 
    from docx import Document


    document = Document()
    document.add_picture('LOGO_SILVERCASTLE.png', width=Inches(1.25))
    h = document.add_heading('Document Title', 0)
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )




    """



    vdoc = "003"
    vdoc_out = "004"
    document = read_template_WORD(f"ENTRADA_ESTABLE_{vdoc}.docx")
    # open connection to Word Document
    #document = docx.Document()



    #document.add_picture('LOGO_2.jpg')
    #h = document.add_heading('', 0)
    #h = document.add_heading('', 0)
    #h = document.add_heading('ACTION PLAN: THEHACKETTAGENT MVP', 0)
    #h = document.add_heading('Document Title', 0)






    records = (
        (3, '000', '0000'),
        (7, '000', '00000'),
        (4, '000', '00000000, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()


    #document.

    document.add_heading('Heading, level 1', level=1)
    #document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.save(f'SALIDA_{vdoc_out}.docx')



main__WORD()
